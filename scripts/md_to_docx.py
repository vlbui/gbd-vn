"""Shared MD -> docx helpers for the manuscript and supplement builders.

The parser is intentionally small: it covers headings (#..####), paragraphs
with inline **bold** / *italic* / `code` / ^sup^, pipe tables, fenced code
blocks, horizontal rules, and blockquotes. It does NOT attempt to be a
general-purpose markdown renderer - only the subset the Lancet manuscripts
actually use.

Lancet formatting conventions applied here:
  - Arial 11pt, 1.5 line spacing everywhere
  - Only ASCII hyphens (em/en/minus replaced with "-")
  - Bold on headings and on leading caption labels like "Figure N." /
    "Table N."; no other bold unless the source has **...** markers

The caller supplies a callback `image_for_caption(text) -> Optional[Path]`
so figure / supplementary-figure paragraphs can have the matching PNG
embedded above the caption line.
"""

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# --- Unicode dash normalization --------------------------------------------
# Map every non-ASCII dash-like char to a plain ASCII hyphen. The task is
# explicit: "0 unicode dashes" and "ASCII dashes only".
_DASH_MAP = {
    "–": "-",  # en dash
    "—": "-",  # em dash
    "−": "-",  # minus sign
    "‐": "-",  # hyphen
    "‑": "-",  # non-breaking hyphen
    "‒": "-",  # figure dash
    "―": "-",  # horizontal bar
}


def normalize_text(s: str) -> str:
    for bad, good in _DASH_MAP.items():
        s = s.replace(bad, good)
    return s


# --- Inline span parser -----------------------------------------------------
# Supported inline markers (parsed in this order to avoid ambiguity):
#   `code`    -> monospace
#   **bold**  -> bold
#   *italic*  -> italic
#   ^sup^     -> superscript
# Anything else is added as a plain run.

_INLINE = re.compile(
    r"(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*|\^[^\s^]+\^)"
)


def _add_inline_runs(paragraph, text, bold=False):
    """Tokenize `text` by inline markers and append appropriately styled
    runs to `paragraph`. `bold` = True forces every run to be bold (used
    for heading paragraphs so that any inline *italic* stays italic+bold).
    """
    text = normalize_text(text)
    parts = _INLINE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Courier New"
            r.bold = bold
        elif part.startswith("^") and part.endswith("^"):
            r = paragraph.add_run(part[1:-1])
            r.font.superscript = True
            r.bold = bold
        elif part.startswith("*") and part.endswith("*"):
            r = paragraph.add_run(part[1:-1])
            r.italic = True
            r.bold = bold
        else:
            r = paragraph.add_run(part)
            r.bold = bold


# --- Document-wide styling --------------------------------------------------

def _style_paragraph(p, font_size=11, bold_heading=False, space_after=6):
    """Apply Arial / 1.5 line-spacing / size to every run in a paragraph,
    and set paragraph-level spacing. Called after runs are added.
    """
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(space_after)
    for r in p.runs:
        r.font.name = r.font.name or "Arial"
        # Make sure East-Asian fallback also picks Arial (docx quirk)
        r.font.size = Pt(font_size)
        rFonts = r._element.rPr.rFonts if r._element.rPr is not None else None
        if rFonts is None:
            rPr = r._element.get_or_add_rPr()
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), r.font.name or "Arial")
        rFonts.set(qn("w:hAnsi"), r.font.name or "Arial")
        rFonts.set(qn("w:cs"), r.font.name or "Arial")
        rFonts.set(qn("w:eastAsia"), "Arial")
        if bold_heading:
            r.bold = True


# --- Table rendering --------------------------------------------------------

def _render_table(doc, lines):
    """Render a pipe-table block (header row, --- separator, data rows)
    into a docx table with Arial 11pt and bold header. Columns without
    explicit widths default to auto.
    """
    def split_row(line):
        # Strip leading/trailing pipe and split; drop empty edges from
        # "| ... |" form while keeping inner empty cells.
        s = line.strip()
        if s.startswith("|"):
            s = s[1:]
        if s.endswith("|"):
            s = s[:-1]
        return [c.strip() for c in s.split("|")]

    header = split_row(lines[0])
    data_lines = lines[2:]  # skip the --- separator
    ncols = len(header)
    table = doc.add_table(rows=1 + len(data_lines), cols=ncols)
    table.style = "Light Grid"

    for j, cell_text in enumerate(header):
        c = table.rows[0].cells[j]
        p = c.paragraphs[0]
        _add_inline_runs(p, cell_text, bold=True)
        _style_paragraph(p, bold_heading=True, space_after=0)

    for i, row_line in enumerate(data_lines, start=1):
        cells = split_row(row_line)
        for j in range(ncols):
            c = table.rows[i].cells[j]
            p = c.paragraphs[0]
            text = cells[j] if j < len(cells) else ""
            _add_inline_runs(p, text)
            _style_paragraph(p, space_after=0)


# --- Block-level parser -----------------------------------------------------

FIGURE_RE = re.compile(r"^\*\*(?:Figure|Fig\.?)\s*(S?\d+)\b", re.IGNORECASE)


def md_to_docx(md_path: Path, out_path: Path, *, image_lookup=None,
               base_font_size=11):
    """Render `md_path` to `out_path`.

    `image_lookup` is a dict mapping figure labels (e.g. "1", "10", "S3")
    to absolute image paths. When a paragraph begins with **Figure N.** or
    **Figure S-N.**, the matching image is inserted immediately before the
    caption paragraph.
    """
    image_lookup = image_lookup or {}
    doc = Document()

    # Narrow margins (Lancet submission)
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    inserted_images = 0
    seen_figure_labels = set()

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        # Horizontal rule -> skip (docx visual separator not needed)
        if line.strip() == "---":
            i += 1
            continue

        # Fenced code block
        if line.startswith("```"):
            i += 1
            code_buf = []
            while i < len(lines) and not lines[i].startswith("```"):
                code_buf.append(lines[i])
                i += 1
            i += 1  # consume closing fence
            p = doc.add_paragraph()
            r = p.add_run(normalize_text("\n".join(code_buf)))
            r.font.name = "Courier New"
            r.font.size = Pt(base_font_size - 1)
            _style_paragraph(p, font_size=base_font_size - 1, space_after=8)
            continue

        # Markdown table (header with pipes, next line is ---|---)
        if ("|" in line and i + 1 < len(lines)
                and re.match(r"^\s*\|?\s*:?-+", lines[i + 1])):
            block = [line]
            j = i + 1
            while j < len(lines) and "|" in lines[j]:
                block.append(lines[j])
                j += 1
            _render_table(doc, block)
            i = j
            continue

        # Headings
        m_head = re.match(r"^(#{1,6})\s+(.+)$", line)
        if m_head:
            level = len(m_head.group(1))
            text = m_head.group(2).strip()
            # Level 1 -> 18pt, 2 -> 14pt, 3 -> 12pt, 4+ -> 11pt
            sizes = {1: 18, 2: 14, 3: 12, 4: 11, 5: 11, 6: 11}
            p = doc.add_paragraph()
            _add_inline_runs(p, text, bold=True)
            _style_paragraph(p, font_size=sizes.get(level, 11),
                             bold_heading=True, space_after=8)
            i += 1
            continue

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Figure caption -> insert image before the caption paragraph
        m_fig = FIGURE_RE.match(line)
        if m_fig and image_lookup:
            label = m_fig.group(1).upper()
            if label not in seen_figure_labels:
                img_path = image_lookup.get(label)
                if img_path and Path(img_path).exists():
                    img_p = doc.add_paragraph()
                    img_p.paragraph_format.space_after = Pt(4)
                    run = img_p.add_run()
                    run.add_picture(str(img_path), width=Inches(6.2))
                    inserted_images += 1
                    seen_figure_labels.add(label)

        # Regular paragraph (may wrap across lines until a blank)
        buf = [line]
        j = i + 1
        while j < len(lines) and lines[j].strip() and not (
                lines[j].startswith("#")
                or lines[j].startswith("```")
                or lines[j].strip() == "---"
                or ("|" in lines[j] and j + 1 < len(lines)
                    and re.match(r"^\s*\|?\s*:?-+", lines[j + 1]))
        ):
            buf.append(lines[j])
            j += 1
        text = " ".join(s.strip() for s in buf)
        p = doc.add_paragraph()
        _add_inline_runs(p, text)
        _style_paragraph(p, font_size=base_font_size, space_after=6)
        i = j

    doc.save(str(out_path))
    return inserted_images
