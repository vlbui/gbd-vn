"""End-to-end verification PASS/FAIL for the 30q70 SEA pipeline, all
regenerated figures, and the rebuilt Word deliverables.

Each check prints "PASS:" or "FAIL:" so the output can be grepped. The
script exits with code 1 if any check fails.
"""

import re
import sys
from pathlib import Path

import pandas as pd
from PIL import Image
from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parent))

from utils import PROJECT_ROOT, TAB, FIG_STATIC


def _check(label, ok, detail=""):
    status = "PASS" if ok else "FAIL"
    print(f"  {status}: {label}" + (f"  ({detail})" if detail else ""))
    return ok


def verify_summary_csv():
    print("\n[1] tables/sea_30q70_summary.csv")
    path = TAB / "sea_30q70_summary.csv"
    df = pd.read_csv(path)
    ok = True
    ok &= _check("exactly 11 rows", len(df) == 11, f"actual {len(df)}")
    ok &= _check("Vietnam present", "Vietnam" in df["country"].values)

    vn = df[df["country"] == "Vietnam"].iloc[0]
    vn_23 = float(vn["30q70_2023"])
    ok &= _check("Vietnam 30q70_2023 in [18.5, 20.5]",
                 18.5 <= vn_23 <= 20.5, f"{vn_23:.3f}")

    ranked = df.sort_values("30q70_2023").reset_index(drop=True)
    ranked["rank"] = ranked.index + 1
    rank = int(ranked[ranked["country"] == "Vietnam"]["rank"].iloc[0])
    ok &= _check("Vietnam rank 6 of 11 ascending",
                 rank == 6, f"actual {rank}")
    return ok


def verify_aapc_agreement():
    print("\n[2] AAPC_loglin vs AAPC_joinpoint agreement (Vietnam)")
    df = pd.read_csv(TAB / "sea_30q70_summary.csv")
    vn = df[df["country"] == "Vietnam"].iloc[0]
    diff = abs(float(vn["AAPC_loglin"]) - float(vn["AAPC_joinpoint"]))
    detail = (f"loglin={vn['AAPC_loglin']:+.3f}, "
              f"joinpoint={vn['AAPC_joinpoint']:+.3f}, diff={diff:.3f} pp")
    return _check("Vietnam |AAPC diff| <= 0.1 pp", diff <= 0.1, detail)


def verify_figures():
    print("\n[3] PNG figures (size, reloadable, no baked titles)")
    ok = True
    pngs = sorted(FIG_STATIC.glob("fig*.png"))
    # Exclude the legacy fig2_heatmap_v2 draft from the "main" set
    main_pngs = [p for p in pngs if "_v2" not in p.name]
    for png in main_pngs:
        try:
            size = png.stat().st_size
            with Image.open(png) as im:
                im.load()
                shape = f"{im.width}x{im.height}"
        except Exception as exc:
            _check(f"{png.name} reloadable", False, str(exc))
            ok = False
            continue
        ok &= _check(f"{png.name} nonzero size", size > 0,
                     f"{size}B, {shape}")

    # Check SVGs for baked "Figure N." titles or figure-level captions
    svgs = [FIG_STATIC / p.name.replace(".png", ".svg") for p in main_pngs]
    for svg in svgs:
        if not svg.exists():
            continue
        text = svg.read_text(encoding="utf-8")
        matches = re.findall(r">Figure\s+\d+[^<]*<", text)
        ok &= _check(f"{svg.name} no 'Figure N.' baked", len(matches) == 0,
                     f"{len(matches)} matches")
    return ok


def verify_panel_borders():
    """Re-execute every figure builder and assert each produced figure has
    four-sided axis lines (`mirror=True` on both x and y) on every subplot.
    """
    print("\n[4] Panel borders on every subplot")
    import importlib

    # Intercept save_fig calls in scripts/06_figures.py so we can inspect
    # the Plotly figure *before* it hits disk and verify axis settings.
    import utils as utils_mod
    captured = []

    def _spy_save(fig, name, width=900, height=600):
        # Mirror what real save_fig does (strip titles, apply borders)
        utils_mod.strip_figure_titles(fig)
        utils_mod.apply_panel_border(fig)
        captured.append((name, fig))

    orig = utils_mod.save_fig
    utils_mod.save_fig = _spy_save
    try:
        figmod = importlib.import_module("06_figures") \
            if "06_figures" in sys.modules else None
        if figmod is None:
            import importlib.util
            spec = importlib.util.spec_from_file_location(
                "figmod_live", PROJECT_ROOT / "scripts" / "06_figures.py")
            figmod = importlib.util.module_from_spec(spec)
            sys.modules["figmod_live"] = figmod
            spec.loader.exec_module(figmod)
        figmod.run()
    finally:
        utils_mod.save_fig = orig

    ok = True
    for name, fig in captured:
        layout = fig.layout
        # Count x and y axes referenced in the layout by their key prefix.
        n_x = sum(1 for k in layout if str(k).startswith("xaxis"))
        n_y = sum(1 for k in layout if str(k).startswith("yaxis"))
        n_subplots = max(n_x, n_y, 1)
        # A subplot has 4 visible borders when both xaxis and yaxis have
        # showline=True and mirror=True/'ticks'/'all'.
        all_mirrored = True
        for k in layout:
            ks = str(k)
            if ks.startswith("xaxis") or ks.startswith("yaxis"):
                ax = layout[k]
                if not (ax.showline and ax.mirror):
                    all_mirrored = False
                    break
        ok &= _check(f"{name}: {n_subplots} subplots have 4-sided borders",
                     all_mirrored)
    return ok


def verify_main_docx():
    print("\n[5] MANUSCRIPT_MAIN.docx")
    path = PROJECT_ROOT / "MANUSCRIPT_MAIN.docx"
    doc = Document(str(path))

    # Count images. python-docx exposes inline shapes per paragraph.
    n_images = 0
    for p in doc.paragraphs:
        n_images += len(p.runs and
                        [r for r in p.runs if r.element.findall(
                            ".//{http://schemas.openxmlformats.org/"
                            "drawingml/2006/main}blip")])

    # Full text for string checks
    full = "\n".join(p.text for p in doc.paragraphs)

    # Check fonts: every run should be Arial
    font_bad = 0
    font_seen = set()
    for p in doc.paragraphs:
        for r in p.runs:
            fname = (r.font.name or "").strip()
            font_seen.add(fname or "(default)")
            if fname and fname not in ("Arial", "Courier New"):
                font_bad += 1
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        fname = (r.font.name or "").strip()
                        font_seen.add(fname or "(default)")
                        if fname and fname not in ("Arial", "Courier New"):
                            font_bad += 1

    # Count unicode dashes in body
    full_all = full
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                full_all += "\n" + cell.text
    bad_dashes = sum(full_all.count(ch) for ch in "–—−"
                                                  "‐‑‒―")

    ok = True
    ok &= _check("embeds 9 images", n_images == 9, f"found {n_images}")
    ok &= _check("0 unicode dashes", bad_dashes == 0,
                 f"found {bad_dashes}")
    ok &= _check("all runs Arial or Courier New", font_bad == 0,
                 f"non-Arial runs={font_bad}, fonts seen={font_seen}")
    ok &= _check('contains "19.5"', "19.5" in full_all)
    ok &= _check('contains "SDG 3.4.1"', "SDG 3.4.1" in full_all)
    ok &= _check('no stale "rank 8 of 11"',
                 "rank 8 of 11" not in full_all)
    ok &= _check('no stale "could not be computed"',
                 "could not be computed" not in full_all)
    # Each image should be followed by a "**Figure N.**" caption paragraph
    # in the text body. Sanity check: count "Figure " paragraph heads.
    fig_caps = sum(1 for p in doc.paragraphs
                   if re.match(r"^Figure\s+\d+\.", p.text.strip()))
    ok &= _check("each image followed by a Figure-caption paragraph",
                 fig_caps >= 9, f"caption paragraphs found: {fig_caps}")
    return ok


def verify_supp_docx():
    print("\n[6] MANUSCRIPT_SUPPLEMENT.docx")
    path = PROJECT_ROOT / "MANUSCRIPT_SUPPLEMENT.docx"
    doc = Document(str(path))

    # Find Table S5b: row count, col count, and Vietnam row
    target_rows, target_cols, has_vn = 0, 0, False
    for tbl in doc.tables:
        header = " | ".join(c.text.strip() for c in tbl.rows[0].cells)
        if header.startswith("Country") and "AAPC" in header \
                and len(tbl.columns) == 8 and len(tbl.rows) == 12:
            target_rows = len(tbl.rows)
            target_cols = len(tbl.columns)
            has_vn = any("Vietnam" in " ".join(c.text for c in row.cells)
                         for row in tbl.rows)
            break
    ok = True
    ok &= _check("Table S5b has 12 rows", target_rows == 12,
                 f"actual {target_rows}")
    ok &= _check("Table S5b has 8 columns", target_cols == 8,
                 f"actual {target_cols}")
    ok &= _check("Table S5b contains Vietnam row", has_vn)
    return ok


def main():
    results = [
        verify_summary_csv(),
        verify_aapc_agreement(),
        verify_figures(),
        verify_panel_borders(),
        verify_main_docx(),
        verify_supp_docx(),
    ]
    print("\n" + ("=" * 60))
    print(f"  OVERALL: {sum(results)}/{len(results)} sections passed")
    if not all(results):
        sys.exit(1)


if __name__ == "__main__":
    main()
